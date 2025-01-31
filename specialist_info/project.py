"""项目信息管理视图"""

import json
from functools import wraps
from random import sample
import docx
from django.shortcuts import render
from django.shortcuts import redirect
from django.http import HttpResponse
from django.core.paginator import Paginator
from django.core.paginator import PageNotAnInteger
from django.core.paginator import EmptyPage
from Specialist.models import *
from Specialist.models import Project
from Specialist.models import ProjectSpecialist
from Specialist.models import UserInfo
from specialist_info.base import base
from .user import check_login

def check_user(fun):
    """检测是否为用户
    用于用户操作页面之前检测，主要为防止直接输入url跳转到无权访问的页面
    - 如果是，则可以直接访问此页面
    - 如果不是，跳转主页
    """
    @wraps(fun)
    def inner(request, *arg, **kwargs):
        level = request.session['login_status']
        if level == '1':
            return fun(request, *arg, **kwargs)
        return redirect('/')
    return inner

@check_login
@check_user
@base
def list_project(request, *arg):
    """项目列表"""
    page = request.GET.get('page')
    username = request.session['username']
    project_list = Project.objects.filter(owner=username)
    paginator = Paginator(project_list, 20)
    try:
        # 用于得到指定页面的内容
        current_page = paginator.page(page)
        # 得到当前页的所有对象列表
        projects = current_page.object_list
    # 请求页码数值不是整数
    except PageNotAnInteger:
        current_page = paginator.page(1)
        projects = current_page.object_list
    # 请求页码数值为空或者在URL参数中没有page
    except EmptyPage:
        # paginator.num_pages返回的是页数
        current_page = paginator.page(paginator.num_pages)
        projects = current_page.object_list
    arg[0]['page'] = current_page
    arg[0]['projects'] = projects
    return render(request, 'project_list.html', arg[0])

@check_login
@check_user
def view_project(request):
    """项目详情"""
    r_id = request.GET.get('id')
    project = Project.objects.get(id=r_id)
    specialist_list = ProjectSpecialist.objects.filter(pid=project)
    response = {}
    response['project'] = project
    response['specialist_list'] = specialist_list
    return render(request, 'project_view.html', response)

@check_login
@check_user
def del_project(request):
    """项目删除"""
    objs = Project.objects.filter(id=request.GET.get('id'))
    if objs:
        for obj in objs:
            obj.delete()
    return redirect('/project/list/')

@check_login
@check_user
def comment_project(request):
    """项目详情"""
    if request.method == 'POST':
        project = Project.objects.get(id=request.POST.get('pid'))
        specialist_list = ProjectSpecialist.objects.filter(pid=project)
        for i, j in request.POST.items():
            if i == 'pid' or i == 'csrfmiddlewaretoken':
                continue
            specialist = Specialist.objects.get(id=int(i))
            obj = specialist_list.get(sid=specialist)
            obj.comment = j
            obj.save()
        response = {}
        response['project'] = project
        response['specialist_list'] = specialist_list
        return redirect('/project/comment?id=' + request.POST.get('pid'))

    r_id = request.GET.get('id')
    project = Project.objects.get(id=r_id)
    specialist_list = ProjectSpecialist.objects.filter(pid=project)
    response = {}
    response['project'] = project
    response['specialist_list'] = specialist_list
    return render(request, 'project_comment.html', response)

@check_login
@check_user
def export(request):
    # 根据 id 获取数据
    r_id = request.GET.get('id')
    project = Project.objects.get(id=r_id)
    specialist_list = ProjectSpecialist.objects.filter(pid=project)
    # 构造 HttpResponse
    response = HttpResponse(content_type='text/docx')  # 指明下载的文件为csv格式
    filename = 'Project.docx'
    response['Content-Disposition'] = 'attachment; filename="{}"'.format(filename)
    # 把内容写入文件
    doc = docx.Document()
    doc.add_heading(project.name, 0)
    doc.add_heading(project.owner.username, 1)
    table = doc.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '姓名'
    hdr_cells[1].text = '性别'
    hdr_cells[2].text = '出生年月'
    hdr_cells[3].text = '类别'
    hdr_cells[4].text = '手机'
    hdr_cells[5].text = '邮箱'
    for specialist in specialist_list:
        row_cells = table.add_row().cells
        row_cells[0].text = specialist.sid.name
        row_cells[1].text = specialist.sid.sex
        row_cells[2].text = specialist.sid.birth.strftime('%Y-%m-%d')
        row_cells[3].text = specialist.sid.category.name
        row_cells[4].text = specialist.sid.phone
        row_cells[5].text = specialist.sid.email
    doc.save(response)
    return response

@check_login
@base
def extractStudent(request, *arg):
    #返回的信息
    response = {}
    #学校数据
    school_list = SchoolList.objects.all()
    #学生数据
    student_list = Studentlist.objects.all()

    for obj in student_list:
        nameAndRank = obj.name + ":" + str(obj.rank)
        response[nameAndRank] = obj.volunteer1

    return HttpResponse(json.dumps(response))

@check_login
@check_user
@base
def extract(request, *arg):
    """学生抽取"""

    #返回的信息
    response = {}
    #学校数据
    school_list = SchoolList.objects.all()
    #学生数据
    student_list = Studentlist.objects.all()
    response['student_list'] = student_list

    return render(request, 'extract_student.html', response)